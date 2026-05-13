"""
Script untuk membuat dokumen Word skripsi lengkap dengan hasil eksperimen
Menggabungkan konten skripsi dengan gambar visualisasi hasil
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import os

# Path
FIGURES_PATH = Path("reports/figures")
OUTPUT_PATH = Path("../SKRIPSI_LENGKAP_DENGAN_GAMBAR.docx")

def create_document():
    """Create complete thesis document with figures"""
    
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # ========== HALAMAN JUDUL ==========
    add_title_page(doc)
    doc.add_page_break()
    
    # ========== ABSTRAK ==========
    add_abstract(doc)
    doc.add_page_break()
    
    # ========== DAFTAR ISI (placeholder) ==========
    add_heading(doc, "DAFTAR ISI", level=0, bold=True, center=True)
    p = doc.add_paragraph()
    p.add_run("(Daftar isi akan diupdate otomatis di Microsoft Word)").italic = True
    doc.add_page_break()
    
    # ========== DAFTAR GAMBAR ==========
    add_list_of_figures(doc)
    doc.add_page_break()
    
    # ========== DAFTAR TABEL ==========
    add_list_of_tables(doc)
    doc.add_page_break()
    
    # ========== BAB 1: PENDAHULUAN ==========
    add_chapter_1(doc)
    doc.add_page_break()
    
    # ========== BAB 2: TINJAUAN PUSTAKA ==========
    add_chapter_2(doc)
    doc.add_page_break()
    
    # ========== BAB 3: METODOLOGI ==========
    add_chapter_3(doc)
    doc.add_page_break()
    
    # ========== BAB 4: HASIL DAN PEMBAHASAN ==========
    add_chapter_4(doc)
    doc.add_page_break()
    
    # ========== BAB 5: PENUTUP ==========
    add_chapter_5(doc)
    doc.add_page_break()
    
    # ========== DAFTAR PUSTAKA ==========
    add_references(doc)
    
    # Save document
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Dokumen berhasil dibuat: {OUTPUT_PATH}")
    print(f"   Ukuran file: {os.path.getsize(OUTPUT_PATH) / 1024 / 1024:.2f} MB")
    return OUTPUT_PATH

def add_title_page(doc):
    """Add title page"""
    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSITAS [NAMA UNIVERSITAS]\nFAKULTAS [NAMA FAKULTAS]\nPROGRAM STUDI [NAMA PRODI]")
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # Logo (placeholder)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[LOGO UNIVERSITAS]")
    run.font.size = Pt(12)
    
    # Spacing
    for _ in range(2):
        doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ANALISIS PERFORMA ALGORITMA MACHINE LEARNING\n")
    run.font.size = Pt(16)
    run.font.bold = True
    run = p.add_run("UNTUK KLASIFIKASI STATUS STUNTING PADA BALITA\n")
    run.font.size = Pt(16)
    run.font.bold = True
    run = p.add_run("MENGGUNAKAN ENSEMBLE LEARNING")
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Spacing
    for _ in range(3):
        doc.add_paragraph()
    
    # SKRIPSI
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SKRIPSI")
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Spacing
    for _ in range(2):
        doc.add_paragraph()
    
    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Disusun oleh:\n[NAMA MAHASISWA]\nNIM: [NIM]")
    run.font.size = Pt(12)
    
    # Spacing
    for _ in range(4):
        doc.add_paragraph()
    
    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[KOTA]\n[TAHUN]")
    run.font.size = Pt(12)

def add_abstract(doc):
    """Add abstract (Intisari and Abstract)"""
    # INTISARI
    add_heading(doc, "INTISARI", level=0, bold=True, center=True)
    
    doc.add_paragraph(
        "Stunting merupakan gangguan pertumbuhan kronis pada balita akibat kurang gizi berkepanjangan yang "
        "berdampak serius terhadap kemampuan kognitif dan produktivitas anak di masa depan. Masalah ini "
        "seringkali terlambat dideteksi karena keterbatasan alat bantu diagnosis yang presisi, sehingga "
        "diperlukan pendekatan teknologi cerdas untuk membantu deteksi dini status gizi balita secara cepat dan akurat."
    )
    
    doc.add_paragraph(
        "Penelitian ini menerapkan metode Ensemble Learning untuk mengklasifikasikan status stunting secara "
        "biner menggunakan lima algoritma, yaitu Decision Tree sebagai baseline, Random Forest, XGBoost, "
        "LightGBM, dan CatBoost. Evaluasi kinerja dilakukan menggunakan Stratified 10-Fold Cross-Validation "
        "pada 39.425 data bersih dengan atribut jenis kelamin, umur, dan tinggi badan."
    )
    
    doc.add_paragraph(
        "Hasil penelitian menunjukkan bahwa seluruh algoritma memperoleh performa sangat tinggi: Decision Tree "
        "99,62%, Random Forest 99,72%, XGBoost 99,67%, LightGBM 99,59%, dan CatBoost 99,67% (akurasi dan F1-Score). "
        "Random Forest menempati peringkat pertama berdasarkan weighted F1-Score, namun uji McNemar's test "
        "menunjukkan bahwa perbedaan antarmodel tidak signifikan secara statistik (α = 0,05)."
    )
    
    doc.add_paragraph(
        "Tinggi badan dan umur terbukti sebagai indikator dominan pada seluruh model. Hasil penelitian ini dapat "
        "dimanfaatkan oleh instansi kesehatan, Puskesmas, dan kader Posyandu sebagai dasar sistem pendukung "
        "keputusan klinis untuk screening awal stunting secara massal dan efisien."
    )
    
    p = doc.add_paragraph()
    run = p.add_run("Kata Kunci: ")
    run.font.bold = True
    p.add_run("Stunting, Machine Learning, Ensemble Learning, Klasifikasi, Random Forest")
    
    doc.add_paragraph()  # Space
    
    # ABSTRACT
    add_heading(doc, "ABSTRACT", level=0, bold=True, center=True)
    
    doc.add_paragraph(
        "Stunting is a chronic growth disorder in toddlers caused by prolonged malnutrition that seriously affects "
        "cognitive ability and future productivity. This condition is often detected late due to the lack of "
        "precise diagnostic tools, necessitating an intelligent technological approach to enable early and accurate "
        "detection of toddler nutritional status."
    )
    
    doc.add_paragraph(
        "This study applies Ensemble Learning methods to classify stunting status in a binary scheme using five "
        "algorithms: Decision Tree as a baseline, Random Forest, XGBoost, LightGBM, and CatBoost. Model performance "
        "was evaluated using Stratified 10-Fold Cross-Validation on 39,425 clean data records with gender, age, "
        "and height as features."
    )
    
    doc.add_paragraph(
        "The results show that all algorithms achieved very high performance: Decision Tree 99.62%, Random Forest "
        "99.72%, XGBoost 99.67%, LightGBM 99.59%, and CatBoost 99.67% (accuracy and F1-Score). Random Forest ranked "
        "first based on weighted F1-Score; however, McNemar's test indicated that the performance differences among "
        "models were not statistically significant (α = 0.05)."
    )
    
    doc.add_paragraph(
        "Height and age were consistently identified as the dominant indicators across all models. These findings "
        "can be utilized by health institutions, community health centers, and Posyandu cadres as a foundation for "
        "clinical decision support systems for mass and efficient early stunting screening."
    )
    
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.font.bold = True
    p.add_run("Stunting, Machine Learning, Ensemble Learning, Classification, Random Forest")

def add_list_of_figures(doc):
    """Add list of figures"""
    add_heading(doc, "DAFTAR GAMBAR", level=0, bold=True, center=True)
    
    figures = [
        "Gambar 3.1  Alur Penelitian",
        "Gambar 3.2  Distribusi Target (4 Kelas Status Gizi)",
        "Gambar 3.3  Ringkasan Informasi Dataset",
        "Gambar 3.4  Deteksi Outlier Menggunakan Metode IQR",
        "Gambar 3.5  Analisis Penghapusan Duplikat",
        "Gambar 3.6  Distribusi Kelas Binary",
        "Gambar 3.7  Pembagian Data Training dan Testing",
        "Gambar 4.1  Distribusi Status Gizi Balita (Original)",
        "Gambar 4.2  Distribusi Umur Balita (Histogram dan Boxplot)",
        "Gambar 4.3  Distribusi Tinggi Badan Balita (Histogram dan Boxplot)",
        "Gambar 4.4  Distribusi Jenis Kelamin",
        "Gambar 4.4  Matriks Korelasi",
        "Gambar 4.5  Distribusi Umur per Kelas",
        "Gambar 4.6  Distribusi Tinggi Badan per Kelas",
        "Gambar 4.7  Hasil Cross-Validation",
        "Gambar 4.8  Dashboard Perbandingan Performa 5 Model (UTAMA)",
        "Gambar 4.9  Perbandingan Confusion Matrix",
        "Gambar 4.10 Perbandingan ROC Curves",
        "Gambar 4.11 Perbandingan Precision-Recall Curves",
        "Gambar 4.12 Perbandingan Metrik Performa",
        "Gambar 4.13 Perbandingan Feature Importance",
        "Gambar 4.14 Distribusi Error Berdasarkan Umur",
        "Gambar 4.15 Distribusi Error Berdasarkan Tinggi Badan",
        "Gambar 5.1  Analisis Trade-off Akurasi vs Waktu Komputasi",
    ]
    
    for i, fig in enumerate(figures, 1):
        p = doc.add_paragraph()
        p.add_run(fig).font.size = Pt(12)
        if i < len(figures):
            p.space_after = Pt(6)

def add_list_of_tables(doc):
    """Add list of tables"""
    add_heading(doc, "DAFTAR TABEL", level=0, bold=True, center=True)
    
    tables = [
        "Tabel 2.1  Penelitian Terkait (Keaslian Penelitian)",
        "Tabel 4.1  Ringkasan Dataset",
        "Tabel 4.2  Hasil Performa Model pada Data Testing",
        "Tabel 4.3  Hasil Uji McNemar's Test",
        "Tabel 4.4  Feature Importance per Model",
        "Tabel 4.5  Ringkasan Error per Model",
    ]
    
    for i, tbl in enumerate(tables, 1):
        p = doc.add_paragraph()
        p.add_run(tbl).font.size = Pt(12)
        if i < len(tables):
            p.space_after = Pt(6)

def add_chapter_1(doc):
    """Add Chapter 1: Pendahuluan"""
    add_heading(doc, "BAB 1", level=0, bold=True, center=True)
    add_heading(doc, "PENDAHULUAN", level=0, bold=True, center=True)
    
    add_heading(doc, "1.1 Latar Belakang", level=1, bold=True)
    
    # Latar belakang lengkap dari revisi_skripsi.md
    doc.add_paragraph(
        "Stunting merupakan kondisi gagal tumbuh kronis pada anak yang ditandai oleh tinggi badan menurut umur "
        "berada di bawah -2 standar deviasi dari median Standar Pertumbuhan Anak WHO [1]. Kondisi ini mencerminkan "
        "malnutrisi kronis yang berlangsung dalam jangka panjang dan dapat mengganggu perkembangan fisik dan "
        "kognitif anak serta meningkatkan risiko penyakit tidak menular pada masa berikutnya [1]. Secara global, "
        "WHO melaporkan bahwa pada tahun 2024 terdapat sekitar 150,2 juta anak usia di bawah lima tahun yang "
        "mengalami stunting atau setara dengan 23,2% populasi balita dunia, sehingga stunting masih menjadi isu "
        "kesehatan masyarakat yang signifikan [2]."
    )
    
    doc.add_paragraph(
        "Di Indonesia, persoalan stunting juga masih membutuhkan perhatian serius. Hasil Survei Status Gizi "
        "Indonesia (SSGI) 2024 yang dirilis Kementerian Kesehatan menunjukkan bahwa prevalensi stunting nasional "
        "telah turun menjadi 19,8% [3]. Meskipun tren ini menunjukkan kemajuan, variasi antardaerah dan "
        "antarkelompok sosial ekonomi masih besar, sehingga kebutuhan terhadap upaya deteksi dini dan intervensi "
        "berbasis data tetap tinggi [3]."
    )
    
    doc.add_paragraph(
        "Perkembangan teknologi komputasi membuka peluang pemanfaatan machine learning untuk membantu klasifikasi "
        "kondisi gizi dan status stunting secara lebih cepat, konsisten, dan terukur [4]. Meta-analisis terbaru "
        "menunjukkan bahwa model machine learning pada data malnutrisi anak memiliki performa moderat hingga baik, "
        "sedangkan berbagai penelitian stunting berbasis data menunjukkan bahwa model berbasis pohon dan ensemble "
        "cenderung memberikan hasil yang kompetitif [4], [10]. Dengan demikian, pendekatan machine learning relevan "
        "untuk dimanfaatkan sebagai pendukung identifikasi dini risiko stunting pada balita."
    )
    
    doc.add_paragraph(
        "Berdasarkan kondisi tersebut, penelitian ini difokuskan pada analisis performa algoritma machine learning, "
        "khususnya lima algoritma berbasis pohon keputusan, untuk klasifikasi status stunting pada balita. Kelima "
        "algoritma tersebut mewakili dua kategori algoritma machine learning, yaitu satu algoritma non-ensemble "
        "sebagai baseline, berupa Decision Tree [13], dan empat algoritma berbasis ensemble learning, yaitu Random "
        "Forest, XGBoost, LightGBM, dan CatBoost [12], [14]."
    )
    
    add_heading(doc, "1.2 Rumusan Masalah", level=1, bold=True)
    doc.add_paragraph(
        "1. Bagaimana performa masing-masing algoritma machine learning (Decision Tree, Random Forest, XGBoost, "
        "LightGBM, dan CatBoost) dalam klasifikasi status stunting pada balita?"
    )
    doc.add_paragraph(
        "2. Algoritma mana yang memberikan performa terbaik berdasarkan metrik evaluasi weighted F1-score?"
    )
    doc.add_paragraph(
        "3. Apakah terdapat perbedaan signifikan secara statistik antarperforma model yang diuji?"
    )
    
    add_heading(doc, "1.3 Tujuan Penelitian", level=1, bold=True)
    doc.add_paragraph(
        "1. Menganalisis performa masing-masing algoritma machine learning untuk klasifikasi status stunting"
    )
    doc.add_paragraph(
        "2. Menentukan algoritma terbaik berdasarkan weighted F1-score"
    )
    doc.add_paragraph(
        "3. Menguji signifikansi perbedaan performa antarmodel menggunakan McNemar's test"
    )
    
    add_heading(doc, "1.4 Manfaat Penelitian", level=1, bold=True)
    doc.add_paragraph(
        "1. Memberikan rekomendasi algoritma terbaik untuk klasifikasi stunting pada balita"
    )
    doc.add_paragraph(
        "2. Menyediakan model siap deploy untuk sistem pendukung keputusan klinis"
    )
    doc.add_paragraph(
        "3. Memberikan kontribusi ilmiah dalam bidang machine learning untuk kesehatan masyarakat"
    )

def add_chapter_2(doc):
    """Add Chapter 2: Tinjauan Pustaka"""
    add_heading(doc, "BAB 2", level=0, bold=True, center=True)
    add_heading(doc, "TINJAUAN PUSTAKA", level=0, bold=True, center=True)
    
    add_heading(doc, "2.1 Penelitian Terkait", level=1, bold=True)
    
    doc.add_paragraph(
        "Penelusuran pustaka pada rentang 2021–2025 menunjukkan bahwa publikasi yang secara eksplisit menggunakan "
        "dataset Kaggle yang sama dengan penelitian ini masih terbatas..."
    )
    
    doc.add_paragraph("[Isi lengkap Bab 2 sesuai revisi_skripsi.md]")
    
    add_heading(doc, "2.2 Landasan Teori", level=1, bold=True)
    add_heading(doc, "2.2.1 Stunting", level=2, bold=True)
    add_heading(doc, "2.2.2 Machine Learning", level=2, bold=True)
    add_heading(doc, "2.2.3 Klasifikasi", level=2, bold=True)
    add_heading(doc, "2.2.4 Ensemble Learning", level=2, bold=True)
    add_heading(doc, "2.2.5 Decision Tree", level=2, bold=True)
    add_heading(doc, "2.2.6 Random Forest", level=2, bold=True)
    add_heading(doc, "2.2.7 XGBoost", level=2, bold=True)
    add_heading(doc, "2.2.8 LightGBM", level=2, bold=True)
    add_heading(doc, "2.2.9 CatBoost", level=2, bold=True)

def add_chapter_3(doc):
    """Add Chapter 3: Metodologi with figures"""
    add_heading(doc, "BAB 3", level=0, bold=True, center=True)
    add_heading(doc, "METODOLOGI PENELITIAN", level=0, bold=True, center=True)
    
    add_heading(doc, "3.1 Alur Penelitian", level=1, bold=True)
    
    doc.add_paragraph(
        "Penelitian ini dilakukan melalui 6 tahap utama sebagai berikut:"
    )
    
    # Add flowchart image
    add_figure(
        doc,
        FIGURES_PATH / "alur_penelitian_placeholder.png",  # Placeholder
        "Gambar 3.1 Alur Penelitian",
        width=6.0,
        add_placeholder=True
    )
    
    add_heading(doc, "3.2 Tahapan Penelitian", level=1, bold=True)
    
    # Tahap 1: Pengumpulan Data
    add_heading(doc, "3.2.1 Pengumpulan Data", level=2, bold=True)
    doc.add_paragraph(
        "Dataset yang digunakan adalah Stunting Toddler Detection yang tersedia di Kaggle dengan "
        "121.000 baris data yang memuat variabel jenis kelamin, umur (bulan), tinggi badan (cm), "
        "dan status gizi balita."
    )
    
    # Add Data Overview figures (DIPECAH - 4 gambar terpisah lebih besar)
    add_figure(
        doc,
        FIGURES_PATH / "01c_distribusi_target.png",
        "Gambar 3.2 Distribusi Target (4 Kelas Status Gizi)",
        width=6.5
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "01d_ringkasan_dataset.png",
        "Gambar 3.3 Ringkasan Informasi Dataset",
        width=6.0
    )
    
    # Tahap 2: Preprocessing
    add_heading(doc, "3.2.2 Pembersihan Data", level=2, bold=True)
    doc.add_paragraph(
        "Tahap preprocessing meliputi deteksi dan penghapusan outlier, penghapusan duplikat, "
        "dan transformasi label ke format biner."
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "02_outlier_detection.png",
        "Gambar 3.3 Deteksi Outlier Menggunakan Metode IQR",
        width=6.0
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "02b_duplicate_analysis.png",
        "Gambar 3.4 Analisis Penghapusan Duplikat",
        width=6.0
    )
    
    # Class Imbalance - gunakan pie chart FIXED (tidak terpotong)
    add_figure(
        doc,
        FIGURES_PATH / "03b_class_imbalance_pie_FIXED.png",
        "Gambar 3.5 Distribusi Kelas Binary (Pie Chart)",
        width=6.5
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "04_train_test_split.png",
        "Gambar 3.6 Pembagian Data Training dan Testing",
        width=6.0
    )

def add_chapter_4(doc):
    """Add Chapter 4: Hasil dan Pembahasan with all figures"""
    add_heading(doc, "BAB 4", level=0, bold=True, center=True)
    add_heading(doc, "HASIL DAN PEMBAHASAN", level=0, bold=True, center=True)
    
    add_heading(doc, "4.1 Analisis Deskriptif Data", level=1, bold=True)
    
    doc.add_paragraph(
        "Dataset yang digunakan dalam penelitian ini terdiri dari 39.425 data bersih setelah "
        "proses preprocessing. Berikut adalah analisis deskriptif terhadap distribusi data."
    )
    
    # Target Distribution - gunakan versi FIXED yang tidak terpotong
    add_figure(
        doc,
        FIGURES_PATH / "target_dist_bar_fixed.png",
        "Gambar 4.1 Distribusi Status Gizi Balita (Bar Chart)",
        width=6.5
    )
    
    # Numerical distributions (DIPECAH - 2 gambar terpisah lebih besar)
    add_figure(
        doc,
        FIGURES_PATH / "numerical_dist_umur_bulan.png",
        "Gambar 4.2 Distribusi Umur Balita (Histogram dan Boxplot)",
        width=6.5
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "numerical_dist_tinggi_badan_cm.png",
        "Gambar 4.3 Distribusi Tinggi Badan Balita (Histogram dan Boxplot)",
        width=6.5
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "categorical_distributions.png",
        "Gambar 4.3 Distribusi Jenis Kelamin",
        width=5.0
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "correlation_heatmap.png",
        "Gambar 4.4 Matriks Korelasi",
        width=5.0
    )
    
    add_heading(doc, "4.2 Hasil Pelatihan Model", level=1, bold=True)
    
    doc.add_paragraph(
        "Lima algoritma machine learning dilatih pada dataset yang sama dengan konfigurasi hyperparameter "
        "yang telah dioptimasi. Evaluasi dilakukan menggunakan stratified 10-fold cross-validation."
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "cv_results_comparison.png",
        "Gambar 4.7 Hasil Cross-Validation",
        width=6.0
    )
    
    add_heading(doc, "4.3 Evaluasi Performa Model", level=1, bold=True)
    
    doc.add_paragraph(
        "Hasil evaluasi menunjukkan bahwa semua model mencapai performa yang sangat tinggi (>98%). "
        "Random Forest menempati peringkat pertama dengan F1-score 99.75%."
    )
    
    # GAMBAR UTAMA - Tabel Performa (FIXED - tidak terpotong)
    add_figure(
        doc,
        FIGURES_PATH / "summary_1_metrics_table_fixed.png",
        "Gambar 4.8 Tabel Perbandingan Performa 5 Model",
        width=6.5,
        important=True
    )
    
    # Dashboard Overview (Alternatif: versi LARGE lengkap)
    add_figure(
        doc,
        FIGURES_PATH / "model_summary_LARGE.png",
        "Gambar 4.9 Dashboard Overview Performa Model (Lengkap)",
        width=6.5,
        important=False
    )
    
    # Add performance table
    add_performance_table(doc)
    
    # Confusion Matrix - Random Forest (terbaik) dan Decision Tree (baseline)
    add_figure(
        doc,
        FIGURES_PATH / "cm_2_rf.png",
        "Gambar 4.10 Confusion Matrix - Random Forest (Model Terbaik)",
        width=6.0
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "cm_1_dt.png",
        "Gambar 4.11 Confusion Matrix - Decision Tree (Baseline)",
        width=6.0
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "summary_2_roc_curves_fixed.png",
        "Gambar 4.12 Perbandingan ROC Curves (FIXED)",
        width=6.5
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "precision_recall_curves.png",
        "Gambar 4.11 Perbandingan Precision-Recall Curves",
        width=6.0
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "metrics_comparison.png",
        "Gambar 4.12 Perbandingan Metrik Performa",
        width=6.0
    )
    
    add_heading(doc, "4.4 Analisis Feature Importance", level=1, bold=True)
    
    doc.add_paragraph(
        "Analisis feature importance menunjukkan bahwa Tinggi Badan adalah fitur paling penting "
        "(63.55% pada Random Forest), diikuti oleh Umur (36.27%), sedangkan Jenis Kelamin memiliki "
        "kontribusi minimal (0.18%)."
    )
    
    # Feature Importance - Random Forest (terbaik)
    add_figure(
        doc,
        FIGURES_PATH / "fi_2_rf.png",
        "Gambar 4.13 Feature Importance - Random Forest",
        width=6.0
    )
    
    add_heading(doc, "4.5 Uji Signifikansi Statistik", level=1, bold=True)
    
    doc.add_paragraph(
        "Uji McNemar's test dilakukan untuk menguji signifikansi perbedaan performa antarmodel. "
        "Hasil menunjukkan bahwa meskipun Random Forest memiliki F1-score tertinggi, perbedaan "
        "dengan model lain tidak signifikan secara statistik (α = 0.05)."
    )
    
    # Add McNemar table
    add_mcnemar_table(doc)
    
    add_heading(doc, "4.6 Analisis Error", level=1, bold=True)
    
    doc.add_paragraph(
        "Analisis error dilakukan untuk memahami karakteristik sampel yang salah klasifikasi. "
        "Hasil menunjukkan bahwa error cenderung terjadi pada kasus borderline, yaitu balita "
        "dengan tinggi badan yang mendekati ambang batas stunting."
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "error_distribution_Umur_bulan.png",
        "Gambar 4.14 Distribusi Error Berdasarkan Umur",
        width=6.5
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "error_distribution_Tinggi_Badan_cm.png",
        "Gambar 4.15 Distribusi Error Berdasarkan Tinggi Badan",
        width=6.5
    )

def add_chapter_5(doc):
    """Add Chapter 5: Penutup"""
    add_heading(doc, "BAB 5", level=0, bold=True, center=True)
    add_heading(doc, "PENUTUP", level=0, bold=True, center=True)
    
    add_heading(doc, "5.1 Kesimpulan", level=1, bold=True)
    
    doc.add_paragraph(
        "1. Kelima algoritma machine learning (Decision Tree, Random Forest, XGBoost, LightGBM, "
        "dan CatBoost) menunjukkan performa sangat tinggi dengan akurasi dan F1-score di atas 98%."
    )
    
    doc.add_paragraph(
        "2. Random Forest mencapai performa terbaik dengan akurasi 99.75%, F1-score 99.75%, "
        "dan ROC-AUC 100% (perfect score)."
    )
    
    doc.add_paragraph(
        "3. Uji McNemar's test menunjukkan bahwa perbedaan performa antarmodel tidak signifikan "
        "secara statistik (α = 0.05), meskipun Random Forest memiliki skor tertinggi."
    )
    
    doc.add_paragraph(
        "4. Tinggi Badan adalah fitur paling penting (63.55%), diikuti Umur (36.27%), sedangkan "
        "Jenis Kelamin memiliki kontribusi minimal (0.18%)."
    )
    
    add_heading(doc, "5.2 Saran", level=1, bold=True)
    
    doc.add_paragraph(
        "1. Random Forest dapat digunakan sebagai model utama untuk sistem deteksi stunting "
        "dengan akurasi sangat tinggi."
    )
    
    doc.add_paragraph(
        "2. Untuk aplikasi mobile atau edge computing, LightGBM atau Decision Tree dapat menjadi "
        "alternatif karena waktu inference yang lebih cepat."
    )
    
    doc.add_paragraph(
        "3. Perlu penelitian lanjutan dengan menambahkan fitur tambahan seperti berat badan, "
        "riwayat kesehatan, dan faktor sosial ekonomi."
    )
    
    doc.add_paragraph(
        "4. Implementasi model dalam sistem klinis perlu dilengkapi dengan mekanisme validasi "
        "manual untuk kasus borderline (confidence rendah)."
    )
    
    # Trade-off analysis
    add_heading(doc, "5.3 Analisis Trade-off", level=1, bold=True)
    
    doc.add_paragraph(
        "Analisis trade-off antara akurasi dan waktu komputasi menunjukkan bahwa Random Forest "
        "memberikan akurasi terbaik dengan waktu inference yang masih acceptable untuk deployment."
    )
    
    add_figure(
        doc,
        FIGURES_PATH / "tradeoff_accuracy_vs_time.png",
        "Gambar 5.1 Analisis Trade-off Akurasi vs Waktu Komputasi",
        width=6.0
    )

def add_references(doc):
    """Add references"""
    add_heading(doc, "DAFTAR PUSTAKA", level=0, bold=True, center=True)
    
    references = [
        "[1] WHO, \"Child growth standards,\" World Health Organization, 2024.",
        "[2] WHO, \"Global nutrition targets 2025: stunting policy brief,\" 2024.",
        "[3] Kementerian Kesehatan RI, \"Hasil Survei Status Gizi Indonesia (SSGI) 2024,\" 2024.",
        "[4] Meta-analysis on ML for child malnutrition, 2024.",
        "[5] R. Y. Pratama and A. Baita, \"Prediksi Stunting pada Anak Balita Menggunakan "
        "Algoritma Extreme Gradient Boosting dan Bayesian Optimization,\" 2025.",
        "[6] D. S. Budianto, \"Comparison of Decision Tree and Support Vector Machine (SVM) "
        "for Stunting Risk Prediction in Toddlers,\" 2025.",
        "[9] O. Pahlevi et al., \"Model Klasifikasi Risiko Stunting Pada Balita Menggunakan "
        "Algoritma CatBoost Classifier,\" 2024.",
        "[10] M. K. Ayele et al., \"Predicting stunting status among under five children in "
        "Ethiopia using ensemble machine learning algorithms,\" 2025.",
        "[11] M. A. H. Hadisuwarno et al., \"Komparasi performa model machine learning algoritma "
        "XGBoost dan Random Forest pada studi kasus mendeteksi stunting,\" 2025.",
        "[12] L. Breiman, \"Random Forests,\" Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.",
        "[13] J. R. Quinlan, \"Induction of decision trees,\" Machine Learning, vol. 1, no. 1, "
        "pp. 81–106, 1986.",
        "[14] T. Chen and C. Guestrin, \"XGBoost: A Scalable Tree Boosting System,\" 2016.",
        "[15] Sklearn documentation on classification metrics.",
        "[16] Kaggle Dataset: Stunting Toddler Detection, 2024.",
        "[17] Q. McNemar, \"Note on the sampling error of the difference between correlated "
        "proportions or percentages,\" Psychometrika, vol. 12, no. 2, pp. 153–157, 1947.",
    ]
    
    for ref in references:
        doc.add_paragraph(ref, style='List Number')

def add_performance_table(doc):
    """Add performance comparison table"""
    p = doc.add_paragraph()
    run = p.add_run("Tabel 4.2 Hasil Performa Model pada Data Testing")
    run.font.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=6, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header
    headers = ['Model', 'Accuracy', 'Precision', 'Recall', 'F1-Score', 'ROC-AUC']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data
    data = [
        ['Random Forest', '99.75%', '99.75%', '99.75%', '99.75%', '100.00%'],
        ['Decision Tree', '99.56%', '99.56%', '99.56%', '99.56%', '99.43%'],
        ['LightGBM', '99.19%', '99.19%', '99.19%', '99.19%', '99.97%'],
        ['CatBoost', '98.88%', '98.88%', '98.88%', '98.88%', '99.93%'],
        ['XGBoost', '98.50%', '98.50%', '98.50%', '98.50%', '99.91%'],
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_paragraph()  # Space after table

def add_mcnemar_table(doc):
    """Add McNemar test results table"""
    p = doc.add_paragraph()
    run = p.add_run("Tabel 4.3 Hasil Uji McNemar's Test")
    run.font.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    headers = ['Pasangan Model', 'p-value', 'Signifikan (α=0.05)']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
    
    # Data
    data = [
        ['Random Forest vs Decision Tree', '< 0.05', 'Ya'],
        ['Random Forest vs XGBoost', '< 0.05', 'Ya'],
        ['Random Forest vs LightGBM', '< 0.05', 'Ya'],
        ['Random Forest vs CatBoost', '< 0.05', 'Ya'],
    ]
    
    for i, row_data in enumerate(data, 1):
        for j, value in enumerate(row_data):
            table.rows[i].cells[j].text = value
    
    doc.add_paragraph()  # Space after table

def add_figure(doc, image_path, caption, width=6.0, important=False, add_placeholder=False):
    """Add figure to document"""
    # Caption before image
    p = doc.add_paragraph()
    run = p.add_run(caption)
    run.font.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if important:
        run.font.color.rgb = RGBColor(255, 0, 0)  # Red for important figures
    
    # Add image
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if add_placeholder or not Path(image_path).exists():
        # Add placeholder text
        run = p.add_run(f"[Gambar: {image_path.name}]\n(Placeholder - gambar akan ditambahkan)")
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    else:
        try:
            run = p.add_run()
            run.add_picture(str(image_path), width=Inches(width))
        except Exception as e:
            print(f"Warning: Could not add image {image_path}: {e}")
            run = p.add_run(f"[Error loading image: {image_path.name}]")
            run.font.italic = True
            run.font.color.rgb = RGBColor(255, 0, 0)
    
    # Space after figure
    doc.add_paragraph()

def add_heading(doc, text, level=1, bold=False, center=False):
    """Add heading to document"""
    if level == 0:
        # Chapter title
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(14)
    elif level == 1:
        # Main heading
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(12)
    else:
        # Sub heading
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.bold = bold
        run.font.size = Pt(12)
    
    return p

if __name__ == "__main__":
    print("\n" + "="*60)
    print("  MEMBUAT DOKUMEN WORD SKRIPSI LENGKAP")
    print("="*60)
    
    try:
        output_file = create_document()
        print(f"\n✅ SUKSES!")
        print(f"\nDokumen tersimpan di:")
        print(f"  {output_file.absolute()}")
        print(f"\nBuka dengan Microsoft Word untuk:")
        print(f"  - Update daftar isi otomatis")
        print(f"  - Tambahkan header/footer dengan nomor halaman")
        print(f"  - Sesuaikan format sesuai panduan skripsi")
        print("\n" + "="*60)
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
